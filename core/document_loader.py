"""文档加载模块

支持 PDF、Markdown、TXT、Word、HTML 格式的文档加载。
"""

import os
from typing import List

from langchain_core.documents import Document
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
)

from config import config


class UnsupportedFormatError(Exception):
    """不支持的文件格式"""
    pass


def _detect_encoding(file_path: str) -> str:
    """检测文件编码"""
    for encoding in ["utf-8", "gbk", "gb2312", "utf-16"]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                f.read()
            return encoding
        except (UnicodeDecodeError, UnicodeError):
            continue
    return "utf-8"


def _load_word(file_path: str) -> List[Document]:
    """加载 Word 文档"""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    content = "\n".join(paragraphs)

    return [Document(
        page_content=content,
        metadata={"source": os.path.basename(file_path)}
    )]


def _load_html(file_path: str) -> List[Document]:
    """加载 HTML 文档"""
    from bs4 import BeautifulSoup

    encoding = _detect_encoding(file_path)
    with open(file_path, "r", encoding=encoding) as f:
        soup = BeautifulSoup(f.read(), "html.parser")

    # 移除 script 和 style 标签
    for tag in soup(["script", "style"]):
        tag.decompose()

    text = soup.get_text(separator="\n", strip=True)

    return [Document(
        page_content=text,
        metadata={"source": os.path.basename(file_path)}
    )]


def load_document(file_path: str) -> List[Document]:
    """加载单个文档

    Args:
        file_path: 文件路径

    Returns:
        List[Document]: LangChain Document 对象列表

    Raises:
        UnsupportedFormatError: 不支持的文件格式
        FileNotFoundError: 文件不存在
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        loader = PyPDFLoader(file_path)
        documents = loader.load()
    elif ext in (".md", ".txt"):
        encoding = _detect_encoding(file_path)
        loader = TextLoader(file_path, encoding=encoding)
        documents = loader.load()
    elif ext in (".docx", ".doc"):
        documents = _load_word(file_path)
    elif ext in (".html", ".htm"):
        documents = _load_html(file_path)
    else:
        raise UnsupportedFormatError(
            f"不支持的文件格式: {ext}。支持的格式: {config.SUPPORTED_EXTENSIONS}"
        )

    # 确保 metadata 包含 source（文件名）
    file_name = os.path.basename(file_path)
    for doc in documents:
        doc.metadata["source"] = file_name

    return documents


def load_documents_from_dir(dir_path: str) -> List[Document]:
    """从目录加载所有支持格式的文档

    Args:
        dir_path: 目录路径

    Returns:
        List[Document]: 所有文档的 Document 对象列表
    """
    if not os.path.isdir(dir_path):
        raise NotADirectoryError(f"目录不存在: {dir_path}")

    all_documents = []

    for root, _, files in os.walk(dir_path):
        for file_name in files:
            ext = os.path.splitext(file_name)[1].lower()
            if ext in config.SUPPORTED_EXTENSIONS:
                file_path = os.path.join(root, file_name)
                try:
                    docs = load_document(file_path)
                    all_documents.extend(docs)
                except Exception as e:
                    print(f"警告: 加载文件 {file_name} 失败: {e}")

    return all_documents
